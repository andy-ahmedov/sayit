from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from pdf_tts_ru.document_extract import extract_document_pages, inspect_document
from pdf_tts_ru.models import TableStrategy
from pdf_tts_ru.pdf_extract import SegmentKind


def test_inspect_document_supports_docx_manual_page_breaks(tmp_path: Path) -> None:
    docx_path = create_sample_docx(tmp_path / "sample.docx")

    inspection = inspect_document(docx_path)

    assert inspection.page_count == 2
    assert [page.table_count for page in inspection.pages] == [1, 0]
    assert inspection.pages[0].char_count > 0


def test_extract_document_pages_docx_inline_embeds_table(tmp_path: Path) -> None:
    docx_path = create_sample_docx(tmp_path / "sample.docx")

    extracted_pages = extract_document_pages(
        docx_path,
        [1, 2],
        table_strategy=TableStrategy.INLINE.value,
    )

    assert len(extracted_pages) == 2
    assert [segment.kind for segment in extracted_pages[0].segments] == [SegmentKind.PROSE]
    assert "Первый абзац." in extracted_pages[0].segments[0].text
    assert "Таблица 1." in extracted_pages[0].segments[0].text
    assert "Вторая страница." in extracted_pages[1].segments[0].text


def test_extract_document_pages_docx_separate_tables(tmp_path: Path) -> None:
    docx_path = create_sample_docx(tmp_path / "sample.docx")

    extracted_pages = extract_document_pages(
        docx_path,
        [1],
        table_strategy=TableStrategy.SEPARATE.value,
    )

    assert [segment.kind for segment in extracted_pages[0].segments] == [
        SegmentKind.PROSE,
        SegmentKind.TABLE,
    ]
    assert "Таблица 1." in extracted_pages[0].segments[1].text


def test_extract_document_pages_markdown_strips_markup(tmp_path: Path) -> None:
    md_path = tmp_path / "notes.md"
    md_path.write_text(
        "# Заголовок\n"
        "- Пункт\n"
        "[ссылка](https://example.com)\n"
        "```python\nprint('skip')\n```\n",
        encoding="utf-8",
    )

    extracted_pages = extract_document_pages(md_path, [1])

    assert extracted_pages[0].segments[0].text == "Заголовок\nПункт\nссылка\n"


def test_extract_document_pages_txt_supports_cp1251(tmp_path: Path) -> None:
    txt_path = tmp_path / "legacy.txt"
    txt_path.write_bytes("Привет из cp1251".encode("cp1251"))

    extracted_pages = extract_document_pages(txt_path, [1])

    assert extracted_pages[0].segments[0].text == "Привет из cp1251"


def test_extract_document_pages_rejects_nonexistent_page_for_single_page_text(tmp_path: Path) -> None:
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text("текст", encoding="utf-8")

    with pytest.raises(ValueError, match="only available page"):
        extract_document_pages(txt_path, [2])


def test_inspect_document_rejects_doc_with_conversion_hint(tmp_path: Path) -> None:
    doc_path = tmp_path / "legacy.doc"
    doc_path.write_bytes(b"not really a doc")

    with pytest.raises(ValueError, match="Convert it to \\.docx first"):
        inspect_document(doc_path)


def create_sample_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("Первый абзац.")

    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Имя"
    table.rows[0].cells[1].text = "Баллы"
    table.rows[1].cells[0].text = "Алиса"
    table.rows[1].cells[1].text = "5"

    break_paragraph = document.add_paragraph()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("Вторая страница.")

    document.save(path)
    return path
